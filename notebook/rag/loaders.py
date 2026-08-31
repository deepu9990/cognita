import os
import re
from abc import ABC, abstractmethod
from typing import List, Optional

from .schemas import DocumentPage, LoadedDocument


class BaseLoader(ABC):
    @abstractmethod
    def load(self, file_path: str, title: Optional[str] = None) -> LoadedDocument:
        pass


class TxtLoader(BaseLoader):
    def load(self, file_path: str, title: Optional[str] = None) -> LoadedDocument:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        filename = os.path.basename(file_path)
        doc_title = title or os.path.splitext(filename)[0]

        # Check for form feed page breaks or split by larger line blocks
        raw_pages = content.split("\f")
        pages: List[DocumentPage] = []

        for idx, page_text in enumerate(raw_pages, start=1):
            cleaned = page_text.strip()
            if cleaned:
                pages.append(
                    DocumentPage(
                        text=cleaned,
                        page_number=idx,
                        section="General",
                        metadata={"source": filename},
                    )
                )

        if not pages and content.strip():
            pages.append(
                DocumentPage(
                    text=content.strip(),
                    page_number=1,
                    section="General",
                    metadata={"source": filename},
                )
            )

        return LoadedDocument(
            filename=filename,
            title=doc_title,
            mime_type="text/plain",
            pages=pages,
            metadata={"file_size": os.path.getsize(file_path)},
        )


class MarkdownLoader(BaseLoader):
    HEADER_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)

    def load(self, file_path: str, title: Optional[str] = None) -> LoadedDocument:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        filename = os.path.basename(file_path)
        doc_title = title or os.path.splitext(filename)[0]

        pages: List[DocumentPage] = []
        lines = content.splitlines()
        current_section = "Introduction"
        current_text: List[str] = []
        page_num = 1

        for line in lines:
            header_match = self.HEADER_PATTERN.match(line)
            if header_match:
                if current_text:
                    section_content = "\n".join(current_text).strip()
                    if section_content:
                        pages.append(
                            DocumentPage(
                                text=section_content,
                                page_number=page_num,
                                section=current_section,
                                metadata={"source": filename},
                            )
                        )
                        page_num += 1
                        current_text = []
                current_section = header_match.group(2).strip()
                current_text.append(line)
            else:
                current_text.append(line)

        if current_text:
            section_content = "\n".join(current_text).strip()
            if section_content:
                pages.append(
                    DocumentPage(
                        text=section_content,
                        page_number=page_num,
                        section=current_section,
                        metadata={"source": filename},
                    )
                )

        return LoadedDocument(
            filename=filename,
            title=doc_title,
            mime_type="text/markdown",
            pages=pages,
            metadata={"file_size": os.path.getsize(file_path)},
        )


class PdfLoader(BaseLoader):
    def load(self, file_path: str, title: Optional[str] = None) -> LoadedDocument:
        filename = os.path.basename(file_path)
        doc_title = title or os.path.splitext(filename)[0]
        pages: List[DocumentPage] = []

        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            for idx, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                cleaned = text.strip()
                if cleaned:
                    # Attempt to detect section header from first line
                    first_line = cleaned.splitlines()[0] if cleaned.splitlines() else ""
                    section = first_line[:80] if len(first_line) < 80 else "Page Content"
                    pages.append(
                        DocumentPage(
                            text=cleaned,
                            page_number=idx,
                            section=section,
                            metadata={"source": filename, "total_pages": len(reader.pages)},
                        )
                    )
        except Exception as err:
            raise RuntimeError(f"Failed to load PDF '{file_path}': {err}")

        return LoadedDocument(
            filename=filename,
            title=doc_title,
            mime_type="application/pdf",
            pages=pages,
            metadata={"file_size": os.path.getsize(file_path), "page_count": len(pages)},
        )


class DocxLoader(BaseLoader):
    def load(self, file_path: str, title: Optional[str] = None) -> LoadedDocument:
        filename = os.path.basename(file_path)
        doc_title = title or os.path.splitext(filename)[0]
        pages: List[DocumentPage] = []

        try:
            import docx

            doc = docx.Document(file_path)
            current_section = "General"
            current_text: List[str] = []
            page_num = 1

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                if para.style.name.startswith("Heading"):
                    if current_text:
                        pages.append(
                            DocumentPage(
                                text="\n".join(current_text),
                                page_number=page_num,
                                section=current_section,
                                metadata={"source": filename},
                            )
                        )
                        page_num += 1
                        current_text = []
                    current_section = text
                else:
                    current_text.append(text)

            # Also parse tables if present
            for table in doc.tables:
                table_lines = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_lines.append(row_text)
                if table_lines:
                    current_text.append("\n".join(table_lines))

            if current_text:
                pages.append(
                    DocumentPage(
                        text="\n".join(current_text),
                        page_number=page_num,
                        section=current_section,
                        metadata={"source": filename},
                    )
                )

        except Exception as err:
            raise RuntimeError(f"Failed to load DOCX '{file_path}': {err}")

        return LoadedDocument(
            filename=filename,
            title=doc_title,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            pages=pages,
            metadata={"file_size": os.path.getsize(file_path)},
        )


def load_document(file_path: str, filename: Optional[str] = None, title: Optional[str] = None) -> LoadedDocument:
    effective_filename = filename or os.path.basename(file_path)
    ext = os.path.splitext(effective_filename)[1].lower()

    if ext == ".pdf":
        return PdfLoader().load(file_path, title=title)
    elif ext in [".docx", ".doc"]:
        return DocxLoader().load(file_path, title=title)
    elif ext in [".md", ".markdown"]:
        return MarkdownLoader().load(file_path, title=title)
    elif ext in [".txt", ".log", ".csv", ".json"]:
        return TxtLoader().load(file_path, title=title)
    else:
        # Fallback to TxtLoader
        return TxtLoader().load(file_path, title=title)

